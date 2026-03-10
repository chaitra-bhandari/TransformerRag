import json
import re
import os
import tempfile
from pathlib import Path
from typing import Dict, List, Any, Optional
from dotenv import load_dotenv
load_dotenv()
from azure.storage.blob import BlobServiceClient

try:
    from docx import Document
except ImportError:
    print("ERROR: pip install python-docx")
    raise

#Azure config
AZURE_CONN_STR     = os.getenv("AZURE_STORAGE_CONNECTION_STRING", "")
TEMPLATE_CONTAINER = os.getenv("BLOB_TEMPLATE_CONTAINER", "order-templates")
OUTPUT_CONTAINER   = os.getenv("BLOB_DESIGN_DOCS_CONTAINER", "order-design-documents")

TEMPLATE_NAMES = {
    "A": "Order A.docx",
    "B": "Order B Template.docx"
}


METADATA_FIELDS = {'value', 'evidence', 'source', 'source_document', 'chunk_id', 'confidence', 'performance', 'unit', 'page'}

#Docx filler class
class DocxFiller:
    """
    Fills Word templates with JSON data (Azure Blob version).
    Includes blocking logic for specific sections.
    """
    
    def __init__(self, json_data: dict, order_type: str = "A"):
        """
        Initialize with JSON data from rag_query.
        
        Args:
            json_data: Dict with extracted parameters
            order_type: 'A' or 'B'
        """
        self.json_data = json_data
        self.order_type = order_type
        self.data = json_data
        self.chunk_lookup = {}
        
        self.blocked_sections = [
            r'Weitere\s+Routineprüfungen', 
            r'Weitere\s+Typprüfungen',       
        ]
        
        #Azure setup
        self.blob_client = BlobServiceClient.from_connection_string(AZURE_CONN_STR)
        self.template_container = self.blob_client.get_container_client(TEMPLATE_CONTAINER)
        self.output_container = self.blob_client.get_container_client(OUTPUT_CONTAINER)
        
        #Create output container if doesn't exist
        try:
            self.output_container.create_container()
        except Exception:
            pass
    
    def normalize(self, text):
        text = re.sub(r'^\d+(\.\d+)*\s*--?\s*', '', text)
        text = re.sub(r'\s*/\s*', ' ', text)
        text = text.lower().strip()
        text = re.sub(r'[^\w\s]', '', text)
        return ' '.join(text.split())
    

    def find_json_key(self, heading):
        norm_heading = self.normalize(heading)
        
        for key in self.data.keys():
            norm_key = self.normalize(key).replace('_', ' ')
            
            if norm_key in norm_heading or norm_heading in norm_key:
                return key
            
            words_h = set(norm_heading.split())
            words_k = set(norm_key.split())
            if words_h and words_k:
                overlap = len(words_h & words_k)
                if overlap > 0 and overlap / max(len(words_h), len(words_k)) >= 0.5:
                    return key
        
        return None
    
    def _apply_unit(self, value, unit):
        """Append unit to a value at every level: scalar, dict leaf, or list item."""
        if not unit or value is None:
            return value
        if isinstance(value, (str, int, float)):
            return f"{value} {unit}"
        if isinstance(value, dict):
            return {k: self._apply_unit(v, unit) for k, v in value.items()}
        if isinstance(value, list):
            return [self._apply_unit(i, unit) for i in value]
        return value

    def strip_metadata(self, data):
        """
        Recursively strip metadata fields from any dict at any nesting level.
        Unit is appended to every leaf value (scalar, dict values, list items).
        """
        if data is None:
            return None

        if isinstance(data, list):
            return [self.strip_metadata(i) for i in data]

        if not isinstance(data, dict):
            return data

        #metadata wrapper: {value, evidence, source, chunk_id, confidence, unit}
        if 'value' in data and any(k in data for k in ('evidence', 'source', 'source_document', 'chunk_id', 'confidence')):
            raw_value   = data.get('value')
            unit        = data.get('unit')
            clean_value = self.strip_metadata(raw_value)
            # Apply unit at every level (scalar, nested dict, or list)
            return self._apply_unit(clean_value, unit)

        #sub_parameters wrapper
        if 'sub_parameters' in data:
            return {k: self.strip_metadata(v) for k, v in data['sub_parameters'].items()}

        #generic dict: recurse, drop metadata/base_term keys
        result = {}
        for k, v in data.items():
            if k in METADATA_FIELDS | {'base_term'}:
                continue
            result[k] = self.strip_metadata(v)

        return result if result else None

    def get_value(self, key):
        if key not in self.data:
            return None, None

        data = self.data[key]
        if not isinstance(data, dict):
            return data, None

        #Get unit before stripping (used for table column headers on nested values)
        unit = data.get('unit') if isinstance(data, dict) else None
        value = self.strip_metadata(data)

        return value, unit
    
    def is_heading(self, para):
        return para.style.name.startswith('Heading') or 'berschrift' in para.style.name
    
    def remove_content_after_heading(self, doc, heading_idx):
        body = doc._element.body
        all_elems = list(body)
        heading_elem = doc.paragraphs[heading_idx]._element
        
        try:
            start_pos = all_elems.index(heading_elem)
        except ValueError:
            return
        
        to_remove = []
        for i in range(start_pos + 1, len(all_elems)):
            elem = all_elems[i]
            
            if elem.tag.endswith('}p'):
                is_heading = False
                for p in doc.paragraphs:
                    if p._element == elem and self.is_heading(p):
                        is_heading = True
                        break
                if is_heading:
                    break
                to_remove.append(elem)
            elif elem.tag.endswith('}tbl'):
                to_remove.append(elem)
        
        for elem in to_remove:
            elem.getparent().remove(elem)
    
    def make_table(self, doc, data):
        if not data:
            return None

        rows = []

        if isinstance(data, dict):
            #Check if any value is itself a non-trivial dict (truly nested table)
            truly_nested = any(
                isinstance(v, dict) and v and not all(k in METADATA_FIELDS for k in v)
                for v in data.values()
            )

            if truly_nested:
                #Multi-column table: one column per sub-key
                cols = set()
                for v in data.values():
                    if isinstance(v, dict):
                        cols.update(k for k in v.keys() if k not in METADATA_FIELDS)

                cols = sorted(cols)
                rows.append(['Parameter'] + cols)

                for k, v in data.items():
                    if isinstance(v, dict):
                        row = [k]
                        for col in cols:
                            val = v.get(col, '-')
                            row.append(str(val) if val is not None else '-')
                        rows.append(row)
            else:
                # Simple Parameter | Value table (handles sub_parameters output)
                rows.append(['Parameter', 'Value'])
                for k, v in data.items():
                    if k in METADATA_FIELDS:
                        continue
                    display_val = '-' if v is None else str(v)
                    rows.append([k, display_val])

        elif isinstance(data, list):
            rows.append(['#', 'Value'])
            for i, item in enumerate(data, 1):
                rows.append([str(i), str(item)])

        if not rows or len(rows) <= 1:
            return None

        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'

        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                table.rows[i].cells[j].text = str(cell_text)
                if i == 0:
                    for run in table.rows[i].cells[j].paragraphs[0].runs:
                        run.bold = True

        return table._element
    
    def make_cooling_table(self, doc, cooling_types):
        """
        Build a table: cooling types + rated power per cooling type.
        Unit is already embedded in values by strip_metadata, so no unit in headers.
        """
        rated_power_data, _ = self.get_value('rated_power')

        if not rated_power_data or not isinstance(rated_power_data, dict):
            rows = [['Cooling Type']]
            for ct in cooling_types:
                rows.append([ct])
        else:
            sample = next(iter(rated_power_data.values()), None)

            if isinstance(sample, dict):
                # Nested: {"ONAN": {"continuous": "500 MVA", "peak": "600 MVA"}}
                all_cols = set()
                for ct in cooling_types:
                    ct_data = rated_power_data.get(ct)
                    if isinstance(ct_data, dict):
                        all_cols.update(ct_data.keys())

                if all_cols:
                    cols = sorted(all_cols)
                    rows = [['Cooling Type'] + cols]
                    for ct in cooling_types:
                        row = [ct]
                        ct_data = rated_power_data.get(ct, {})
                        for col in cols:
                            val = ct_data.get(col) if isinstance(ct_data, dict) else None
                            row.append(str(val) if val is not None else '-')
                        rows.append(row)
                else:
                    rows = [['Cooling Type']]
                    for ct in cooling_types:
                        rows.append([ct])
            else:
                # Flat: {"ONAN": "500 MVA", "ONAF": "750 MVA"}
                rows = [['Cooling Type', 'Rated Power']]
                for ct in cooling_types:
                    val = rated_power_data.get(ct, '-')
                    rows.append([ct, str(val) if val is not None else '-'])

        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'
        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                table.rows[i].cells[j].text = str(cell_text)
                if i == 0:
                    for run in table.rows[i].cells[j].paragraphs[0].runs:
                        run.bold = True

        return table._element
    
    def _parse_chunk_id_string(self, chunk_id_str):
        """
        Parse chunk_id string like 'TenneT_715MVA_SPE.01.456-G-VK715_text_37'
        into (doc_name, page).

        Format: {ProjectName}_{DocName}_{text|table}_{number}
        """
        if not chunk_id_str:
            return '-', '-'

        # First try pkl lookup (most accurate)
        if chunk_id_str in self.chunk_lookup:
            info = self.chunk_lookup[chunk_id_str]
            doc_name = Path(info['source']).name if info.get('source') else chunk_id_str
            page = str(info['page']) if info.get('page') else '-'
            return doc_name, page

        # Fallback: parse the string
        # Strip trailing _{text|table}_{number}
        match = re.match(r'^(.+?)_(text|table)_(\d+)$', chunk_id_str)
        if match:
            base = match.group(1)   # e.g. TenneT_715MVA_SPE.01.456-G-VK715
            page = match.group(3)   # e.g. 37

            # Extract doc name: find first segment containing a dot or dash
            # (document IDs like SPE.01.456-G-VK715 have dots/dashes)
            parts = base.split('_')
            doc_parts = []
            for part in parts:
                if '.' in part or '-' in part:
                    doc_parts = parts[parts.index(part):]
                    break

            doc_name = '_'.join(doc_parts) if doc_parts else base
            return doc_name, page

        return chunk_id_str, '-'

    def _get_source_info(self, key):
        """
        Get (source_document, page) for a given JSON key.
        Extracts source_document and page directly from JSON structure.
        """
        raw = self.data.get(key, {})

        def find_refs(d):
            if not isinstance(d, dict):
                return None, None

            #Check for source_document and page fields
            if 'source_document' in d or 'page' in d:
                source_doc = d.get('source_document', '-')
                page       = d.get('page')
                
                #Handle page conversion safely
                page_str = '-'
                if page is not None:
                    try:
                        page_str = str(int(page))
                    except (ValueError, TypeError):
                        #If page is 'none' or any other invalid value, use '-'
                        page_str = '-' if str(page).lower() in ['none', 'null', 'n/a'] else str(page)
                
                return source_doc or '-', page_str

            # Recurse into nested dicts
            for v in d.values():
                if isinstance(v, dict):
                    doc, pg = find_refs(v)
                    if doc and doc != '-':
                        return doc, pg

            return None, None

        source_doc, page = find_refs(raw)
        return (source_doc or '-'), (page or '-')
     
    #Process document with blocking logic for specific sections 
    def process(self, doc):
        #Find all headings and match to JSON keys
        actions = []
        for i, para in enumerate(doc.paragraphs):
            if self.is_heading(para):
                heading_text = para.text.strip()
                
                #Skip sections that are in blocked_sections list (leave them blank)
                skip_section = False
                for pattern in self.blocked_sections:
                    if re.search(pattern, heading_text, re.IGNORECASE):
                        skip_section = True
                        print(f"BLOCKED: {heading_text}")
                        break
                
                if skip_section:
                    continue
                
                key = self.find_json_key(heading_text)
                if key:
                    value, unit = self.get_value(key)
                    if value is not None:
                        actions.append((i, key, value, unit, heading_text))
                        print(f"   ✓ {heading_text} → {key}")
        
        print(f"Matched: {len(actions)} sections\n")

        #Process in reverse order
        for idx, key, value, unit, heading in reversed(actions):
            self.remove_content_after_heading(doc, idx)

            heading_elem = doc.paragraphs[idx]._element
            parent = heading_elem.getparent()
            pos = parent.index(heading_elem)

            new_elems = []
            unit_str = f" {unit}" if unit else ""

            #Frequency special case
            if key == 'frequency' and isinstance(value, (int, float)):
                p = doc.add_paragraph(f"Frequency: {value}{unit_str}")
                new_elems.append(p._element)

            #Cooling type - special combined table
            elif key == 'cooling_type' and isinstance(value, list):
                table = self.make_cooling_table(doc, value)
                if table:
                    new_elems.append(table)

            #String
            elif isinstance(value, str):
                for line in value.split('\n'):
                    if line.strip():
                        p = doc.add_paragraph(line)
                        new_elems.append(p._element)

            #Number
            elif isinstance(value, (int, float)):
                p = doc.add_paragraph(f"{value}{unit_str}")
                new_elems.append(p._element)

            #List of simple values
            elif isinstance(value, list):
                if all(isinstance(x, (str, int, float)) for x in value):
                    for i, item in enumerate(value, 1):
                        p = doc.add_paragraph(f"{i}. {item}")
                        new_elems.append(p._element)
                else:
                    table = self.make_table(doc, value)
                    if table:
                        new_elems.append(table)

            #Dict → table
            elif isinstance(value, dict):
                table = self.make_table(doc, value)
                if table:
                    new_elems.append(table)

            #Spacing
            new_elems.append(doc.add_paragraph("")._element)

            #Insert elements after heading
            for elem in reversed(new_elems):
                parent.insert(pos + 1, elem)

        #Only parameters that were actually filled, with source_document + page
        filled_keys = [key for (_, key, _, _, _) in actions]

        if filled_keys:
            doc.add_heading('References', level=1)

            #Build reference rows
            ref_rows = [['Parameter', 'Source Document', 'Page']]
            for key in filled_keys:
                source_doc, page = self._get_source_info(key)
                if source_doc and source_doc != '-':
                    page_str = str(page) if page else '-'
                    ref_rows.append([key.replace('_', ' ').title(), source_doc, page_str])

            if len(ref_rows) > 1:  
                ref_table = doc.add_table(rows=len(ref_rows), cols=3)
                ref_table.style = 'Table Grid'

                for i, row_data in enumerate(ref_rows):
                    for j, cell_text in enumerate(row_data):
                        ref_table.rows[i].cells[j].text = str(cell_text)
                        if i == 0:
                            for run in ref_table.rows[i].cells[j].paragraphs[0].runs:
                                run.bold = True
    
    #Download template from Azure blob
    def download_template(self, temp_dir: str) -> str:
        template_name = TEMPLATE_NAMES[self.order_type]
        local_path = Path(temp_dir) / template_name
        
        data = (self.template_container
                    .get_blob_client(template_name)
                    .download_blob()
                    .readall())
        
        with open(local_path, "wb") as f:
            f.write(data)
        
        print(f"Template downloaded")
        return str(local_path)
    
    #Upload filled document to Azure blob
    def upload_output(self, local_path: str, project_name: str) -> str:
        blob_name = f"{project_name}/Order_{self.order_type}.docx"
        
        with open(local_path, "rb") as f:
            self.output_container.upload_blob(name=blob_name, data=f, overwrite=True)
        
        print(f" Uploaded to {OUTPUT_CONTAINER}/{blob_name}")
        return blob_name
    
    #Full pipeline: Download → Fill → Upload
    def run(self, project_name: str) -> str:
        print(f"\n{'='*60}")
        print(f"FILLING ORDER {self.order_type}")
        print(f"{'='*60}")
        print(f"Project: {project_name}")
        print(f"Sections: {len(self.data)}\n")
        
        with tempfile.TemporaryDirectory() as temp_dir:
            #Download template
            template_path = self.download_template(temp_dir)
            
            #Load and process
            doc = Document(template_path)
            self.process(doc)
            
            #Save locally
            output_path = str(Path(temp_dir) / f"Order_{self.order_type}_filled.docx")
            doc.save(output_path)
            
            #Upload to blob
            blob_path = self.upload_output(output_path, project_name)
            
            print(f"Order {self.order_type} complete!\n")
            return blob_path